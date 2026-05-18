import io

import docx as _docx
import pypdf
import pytest

from app.parsers import (
    UnsupportedFormatError,
    parse,
    parse_docx,
    parse_markdown,
    parse_pdf,
    parse_text,
)


def _make_pdf(text: str) -> bytes:
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(width=200, height=200)
    page.merge_page(pypdf.PageObject.create_blank_page(width=200, height=200))
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_pdf_with_text(text: str) -> bytes:
    """Minimal valid PDF with embedded text via direct content stream."""
    content = b"BT /F1 12 Tf 50 700 Td (" + text.encode("latin-1", errors="replace") + b") Tj ET"
    resources = b"<< /Font << /F1 << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> >> >>"
    page_obj = (
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]"
        b" /Contents 4 0 R /Resources " + resources + b" >>"
    )
    pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj " + page_obj + b" endobj\n"
        b"4 0 obj << /Length "
        + str(len(content)).encode()
        + b" >>\nstream\n"
        + content
        + b"\nendstream endobj\n"
        b"xref\n0 5\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000274 00000 n \n"
        b"trailer << /Size 5 /Root 1 0 R >>\nstartxref\n400\n%%EOF"
    )
    return pdf


def _make_docx(text: str) -> bytes:
    doc = _docx.Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parse_pdf_returns_str():
    payload = _make_pdf("ignored")
    result = parse_pdf(payload)
    assert isinstance(result, str)


def test_parse_docx_returns_text():
    payload = _make_docx("Hello from DOCX")
    result = parse_docx(payload)
    assert "Hello from DOCX" in result


def test_parse_docx_multiple_paragraphs():
    doc = _docx.Document()
    doc.add_paragraph("First")
    doc.add_paragraph("Second")
    buf = io.BytesIO()
    doc.save(buf)
    result = parse_docx(buf.getvalue())
    assert "First" in result
    assert "Second" in result


def test_parse_dispatch_pdf():
    payload = _make_pdf("test")
    result = parse("file.pdf", payload)
    assert isinstance(result, str)


def test_parse_dispatch_docx():
    payload = _make_docx("test content")
    result = parse("report.docx", payload)
    assert "test content" in result


def test_parse_text_ascii():
    assert parse_text(b"hello world") == "hello world"


def test_parse_text_utf8_cyrillic():
    assert parse_text("привет мир".encode()) == "привет мир"


def test_parse_text_invalid_bytes_replaced():
    result = parse_text(b"\xff\xfe abc")
    assert "abc" in result
    assert "�" in result


def test_parse_text_empty():
    assert parse_text(b"") == ""


def test_parse_markdown_no_frontmatter_unchanged():
    src = "# title\n\nbody text\n"
    assert parse_markdown(src.encode()) == src


def test_parse_markdown_strips_frontmatter():
    src = b"---\ntitle: x\nauthor: me\n---\n# body\n"
    assert parse_markdown(src) == "# body\n"


def test_parse_markdown_strips_crlf_frontmatter():
    src = b"---\r\ntitle: x\r\n---\r\n# body\r\n"
    assert parse_markdown(src) == "# body\r\n"


def test_parse_markdown_strips_frontmatter_closed_with_dots():
    src = b"---\ntitle: x\n...\n# body\n"
    assert parse_markdown(src) == "# body\n"


def test_parse_markdown_keeps_horizontal_rule_in_body():
    src = "# title\n\n---\n\nbody\n"
    assert parse_markdown(src.encode()) == src


def test_parse_markdown_strips_empty_frontmatter():
    src = b"---\n---\n# body\n"
    assert parse_markdown(src) == "# body\n"


def test_parse_markdown_unclosed_frontmatter_left_alone():
    src = "---\ntitle: x\nno close here\n"
    assert parse_markdown(src.encode()) == src


def test_parse_dispatch_txt():
    assert parse("notes.txt", b"plain") == "plain"


def test_parse_dispatch_md():
    assert parse("readme.md", b"---\nt: x\n---\n# h\n") == "# h\n"


def test_parse_dispatch_markdown_extension():
    assert parse("doc.markdown", b"# h\n") == "# h\n"


def test_parse_dispatch_uppercase_extension():
    assert parse("README.MD", b"# h\n") == "# h\n"


def test_parse_unknown_extension_raises():
    with pytest.raises(UnsupportedFormatError, match="unsupported file extension"):
        parse("blob.bin", b"\x00\x01")


def test_parse_empty_filename_raises():
    with pytest.raises(UnsupportedFormatError, match="<none>"):
        parse("", b"data")


def test_parse_filename_without_extension_raises():
    with pytest.raises(UnsupportedFormatError, match="<none>"):
        parse("noext", b"data")
